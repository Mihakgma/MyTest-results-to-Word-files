from ParserModule import MyTestResultParser
from DateCheckerFile import DateChecker
from pandas import DataFrame
import os
import docx

#from VariableValuesMarks import all_variables_marks
#print(all_variables_marks['PC_name'][1])


previous_wd_path = os.getcwd()
# тестируем класс MyTestResultParser

# прерка корректности работы проверки формата файла для парсинга!!!
testObj = MyTestResultParser('ksdhfkjsdlfj.txt', 'ksdhfkjsdlfj.txt')

print(testObj.check_format()[0])
print(testObj.check_format()[1])

# проверяем класс проверки дат из строки в русском формате (ДД.ММ.ГГГГ)!!!


test_dates_lst = [
    'kdfjngkds',  # не дата
    '01.01.2000',
    '30.02.2001',
    '20.10.2021',
    '  09.10.2022  ',
    '01-7-1990',  # неправильный формат (разделители)
    '01.02.2019 13:12:00',
    '01.02.2019 13:12:60', # дата с ошибкой во времени
    '01.17.2019 03:33:60'  # дата с ошибкой в месяце
]

for date in test_dates_lst:
    if len(date.strip()) > 10:
        result_date_check = DateChecker(date_to_check=date).check_date(date_format='long')
    else:
        result_date_check = DateChecker(date_to_check=date).check_date()
    print(result_date_check)

# пробуем подгружать файл с результатами тестирования из MyTest!!!
file_full_path = r'D:\ФБУЗ_ЦГиЭКО\ТЕСТЫ_ОГВиА\MyTestSavePythonApplication\файл для парсинга\MyTestStudent_Result 19.01.2022.txt'

print(file_full_path)
testObjRealPath = MyTestResultParser(file_full_path, file_full_path)
print(testObjRealPath.check_format()[0])
print(testObjRealPath.check_format()[1])
# проверяем правильность считывания содержимого файла!!!
file_cont = testObjRealPath.read_file()
#print(file_cont)
print('\n', testObjRealPath.get_all_attributes())
# проверяем правильность считывания содержимого файла!!!
testObjRealPath.content_slicer('manual', 'Текст для разделения!!!')
#testObjRealPath.content_slicer()
# пробник слайсера
slice_results = testObjRealPath.content_slicer()
#print([elem for elem in sorted(file_cont.split(),key=len) if 'n' in elem])
#print(f'\n\nСодержимое оставшегося текста, в котором не было найдено результатов тестирования: \n{slice_results[1]}')
print(slice_results[0])
# сохраняем результаты тестирования в виде эксель-файла!
#df_temp = DataFrame(slice_results[0]).T # транспонированный ДФ!!!
df_temp = DataFrame(slice_results[0])
print(df_temp.head())
excel_filename = 'test_output_MyTestProgramSliced.xlsx'
df_temp.to_excel(excel_filename, index=False)
print(os.getcwd())
fileLstCWD = os.listdir()
if excel_filename in fileLstCWD:
    print('Указанный эксель-файл успешно сохранен в текущую рабочую директорию.')
else:

    print('Указанный эксель-файл НЕ БЫЛ СОХРАНЕН в текущую рабочую директорию!?')

# пробник на создание директории на рабочем столе и поддиректорий с датами тестирования
#today_date = dt.today().strftime('%d.%m.%Y')
#
desktop_path = str(os.environ['USERPROFILE'] + '\Desktop') + f'\\ВЫГРУЗКА_РЕЗУЛЬТАТОВ_ИЗ_MY_TEST\\'

def filename_purify(filename: str):
    filename_purified = filename[:]

    for elem in [
        ':',
        '/',
        '\\',
        '*',
        '?',
        '"',
        '<',
        '>',
        '|'
    ]:
        filename_purified = filename_purified.replace(elem, '_')

    return filename_purified

#try:
for colname in list(df_temp):
    current_column_contents = df_temp[colname].to_list()
    current_date = current_column_contents[2][:10]
    date_result_current = DateChecker(date_to_check=current_date).check_date()
    if date_result_current['date_ok']:
        current_date_full = current_date + '\\'
        try:
            dir_fullpath = desktop_path+current_date_full
            os.makedirs(dir_fullpath)
        except:
            print(f'Папка <ВЫГРУЗКА_РЕЗУЛЬТАТОВ_ИЗ_MY_TEST\{current_date}> была создана ранее на Вашем рабочем столе!')
        current_censored_result = current_column_contents[1]

        os.chdir(dir_fullpath)

        # txt-file
        #txt_filename = colname + '.txt'
        #f = open(filename_purify(txt_filename), 'w+')
        #f.write(current_censored_result)
        #f.close()

        # docx-file
        os.chdir(previous_wd_path)
        doc = docx.Document('пустой шаблон.docx')
        row = 0
        for paragraph in doc.paragraphs:
            if row == 0:
                paragraph.add_run(current_censored_result)
            row += 1

        os.chdir(dir_fullpath)
        doc.save(f'{filename_purify(colname)}.docx')

#except:
#    print('Папка с названием ВЫГРУЗКА_РЕЗУЛЬТАТОВ_ИЗ_MY_TEST была ранее создана на Вашем рабочем столе!')
# возврат в изначальную рабочую директорию!!!
os.chdir(previous_wd_path)

input()