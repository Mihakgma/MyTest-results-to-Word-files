from ParserModule import MyTestResultParser
from DateCheckerFile import DateChecker
from pandas import DataFrame
from pandas import concat as pd_concat
import os
import docx



# сохраняем текущую рабочую директорию в виде строки
previous_wd_path = os.getcwd()

# ДОБАВЛЕНА ВОЗМОЖНОСТЬ ВЫБОРА ФАЙЛА,
# НО ИЗ СПЕЦИАЛЬНО ПРЕДНАЗНАЧЕННОЙ ДЛЯ ЭТОГО ПАПКИ!!!

#############################
# ФУНКЦИИ
############################
def select_filename_to_parse():
    """
    Функция выбора наименования файла для
    извлечения результатов тестирования
    """
    filelist_cwd = os.listdir('.\\папка для результатов тестирования\\')
    [print(f'{file_index}: <{filelist_cwd[file_index]}>\n') for file_index in range(len(sorted(filelist_cwd)))]

    counter = 0
    while counter < 6:
        counter += 1
        try:
            chosen_list_index = int(input('Введите номер файла из '
                                          'предложенного выше перечня файлов '
                                          'для извлечения результатов тестирования: \n'))
            chosen_file_name = filelist_cwd[chosen_list_index]
            print(f'Вы выбрали следующий файл для извлечения результатов тестирования: {chosen_file_name}')
            answer = input('Вы уверены? (Д/Н)')
            if 'н' in answer.lower() or 'n' in answer.lower():
                pass
            else: # пользователь подтверждает правильность выбранного им файла из предложенного перечня
                return chosen_file_name
        except ValueError:
            print('Что-то пошло не так...')
            print('Возможно Вы ввели не цифру...')
            print('Пожалуйста, повторите ввод.')
        except IndexError:
            print('Что-то пошло не так...')
            print('Возможно, Вы ввели цифру не из предложенного перечня...')
            print('Пожалуйста, повторите ввод.')

        # в случае ошибки функция возвращает пустую строку
        return ''


def filename_purify(filename: str):
    """
    Очищает строку для наименования файла
    от недопустимых символов
    """
    filename_purified = filename[:]
    for elem in [
        ':',
        '/',
        '\\',
        '*',
        '?',
        '"',
        '<',
        '>',
        '|'
    ]:
        filename_purified = filename_purified.replace(elem, '_')
    return filename_purified


def save_df_excel(df, dir_path, today_str: str):
    excel_filename = f'test_output_MyTestProgramSliced_{today_str}.xlsx'
    # ТРАНСПОНИРУЕМ ФРЕЙМ ДАННЫХ!!!
    df = df.T
    df.columns = [
        'текст весь',
        'текст цензурированный',
        'время окончания теста',
        'имя ПК',
        'имя пользователя',
        'ФИО',
        'Название теста',
        'Расположение файла с тестом',
        'CRC файла',
        'всего заданий в тесте',
        'выполнено заданий',
        'из них правильно',
        'из них ошибок',
        'результатоивность',
        'использовано подсказок',
        'набрано баллов',
        'оценка',
        'маска результата',
        'маска времени обдумывания (в секундах)',
        'маска ответов',
        'маска штрафов за подсказку',
        'время начала',
        'время завершения',
        'продолжительность'
    ]
    df.to_excel(dir_path+excel_filename, index=True)
    filelist_cwd = os.listdir(dir_path)
    if excel_filename in filelist_cwd:
        print('Эксель-файл с результатами тестирования УСПЕШНО СОХРАНЕН в директорию для выгрузки данных.')
    else:
        print('Эксель-файл с результатами тестирования НЕ БЫЛ СОХРАНЕН в директорию для выгрузки данных!?')


def get_today_with_time_str():
    date_obj = DateChecker('')
    today_time_str = date_obj.get_today_date('medium')
    return today_time_str


def file_in_directory(filename: str, dir_path: str):
    """
    Проверяет находится ли файл с указанным
    названием в директории с указанным путем.
    Возвращает:
    Да - True
    Нет - False
    """
    checking_directory_filelist = os.listdir(dir_path)
    if filename in checking_directory_filelist:
        return True
    else:
        return False


def main():
    #file_full_path = '.\\папка для результатов тестирования\\MyTestStudent_Result 19.01.2022.txt'
    file_full_path = '.\\папка для результатов тестирования\\' + select_filename_to_parse()
    print(file_full_path)
    testObjRealPath = MyTestResultParser(file_full_path, file_full_path)

    # слайсер содержимого файла
    slice_results = testObjRealPath.content_slicer()
    # сохраняем результаты тестирования в виде эксель-файла!
    counter = 0
    for dict_current in slice_results[0]:
        counter += 1
        df_current = DataFrame(dict_current)
        # 1-ый ДФ берется за основу, к которому затем добавляются остальные по колонкам (т.е. вниз)
        if counter == 1:
            df_temp = df_current.copy()
        else:
            df_list = []
            df_list.append(df_temp)
            df_list.append(df_current)
            df_temp = pd_concat(df_list, axis=1)
    #print(df_temp.head())
    # пробник на создание директории на рабочем столе и поддиректорий с датами тестирования

    desktop_import_dir_path = str(os.environ['USERPROFILE'] + '\Desktop') + f'\\ВЫГРУЗКА_РЕЗУЛЬТАТОВ_ИЗ_MY_TEST\\'

    # проверяем была ли ранее создана папка на рабочем столе с данным именем
    if not(os.path.exists(desktop_import_dir_path)):
        os.makedirs(desktop_import_dir_path)

    # сохраняем результаты в директорию с выгрузкой
    today_str = filename_purify(get_today_with_time_str())
    save_df_excel(df_temp, desktop_import_dir_path, today_str)

    # сюда будут добавляться вновь добавленные файлы
    # но только те, которых ранее не было в папках!!!
    new_files_list = ''
    # строка для перезаписанных файлов - были ранее в папках!!!
    rerecorded_files_list = ''
    for colname in list(df_temp):
        current_column_contents = df_temp[colname].to_list()
        current_date = current_column_contents[2][:10]
        date_result_current = DateChecker(date_to_check=current_date).check_date()
        current_year = str(date_result_current['year_num'])
        current_month = date_result_current['month_str'][:3]
        current_month_num = str(date_result_current['month_num'])
        if date_result_current['date_ok']:
            current_date_full = '\\'+current_year+'\\'+current_month_num+'-'+current_month+'-'+current_year+'\\'+current_date+'\\'
            try:
                dir_fullpath = desktop_import_dir_path + current_date_full
                os.makedirs(dir_fullpath)
            except:
                pass
                # print(f'Папка <ВЫГРУЗКА_РЕЗУЛЬТАТОВ_ИЗ_MY_TEST\{current_date_full[1:]}> '
                #      f'была создана ранее на Вашем рабочем столе!')
            current_censored_result = current_column_contents[1]

            os.chdir(dir_fullpath)

            # docx-file
            os.chdir(previous_wd_path)
            doc = docx.Document('.\\шаблон docx-файла\\пустой шаблон.docx')
            row = 0
            for paragraph in doc.paragraphs:
                if row == 0:
                    paragraph.add_run(current_censored_result)
                row += 1

            os.chdir(dir_fullpath)
            docx_filename = f'{filename_purify(colname)}.docx'

            # проверка на наличие файла стаким именем в текущей директории
            if file_in_directory(docx_filename, dir_fullpath):
                rerecorded_files_list += dir_fullpath + docx_filename + '\n'
            else: # если файла в директории РАНЕЕ НЕ БЫЛО
                new_files_list += dir_fullpath + docx_filename + '\n'

            doc.save(docx_filename)


    # меняем рабочую директорию на корневую папку с выгрузкой
    os.chdir(desktop_import_dir_path)

    # ВНОВЬ СОЗДАННЫЕ ФАЙЛЫ
    txt_results_filename = f'НОВЫЕ_файлы_с_результатами_в_текущей_директории_({today_str})' + '.txt'
    txt_file = open(txt_results_filename, 'w+')
    txt_file.write(new_files_list) # записываем в файл строку с результатами
    txt_file.close()

    # ПЕРЕЗАПИСАННЫЕ ФАЙЛЫ
    txt_results_filename = f'ПЕРЕЗАПИСАННЫЕ_файлы_с_результатами_в_текущей_директории_({today_str})' + '.txt'
    txt_file = open(txt_results_filename, 'w+')
    txt_file.write(rerecorded_files_list) # записываем в файл строку с результатами
    txt_file.close()

    input()


#############################
# ОСНОВНОЙ БЛОК КОДА
############################
if __name__ == '__main__':
    main()