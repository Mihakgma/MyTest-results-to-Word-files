#!/usr/bin/env python
# coding: utf-8

# In[8]:



"""
ТЕСТОВАЯ ВЕРСИЯ!

Данная программа принимает на вход экселевский файл 
(в формате *.xls или *.xlsx - на выбор),
парсит лист (ручной выбор названия листа файла пользователем)
данного файла в пандовский датафрэйм,
всех пассажиров согласно списка разделяет на детей и взрослых.

Спрашивает у пользователя откуда рейс (напр. Москва).

Выводит на экран информацию по:
1) общему количеству пассажиров;
2) количеству совершеннолетних пассажиров;
3) количеству несовершеннолетних пассажиров;
4) количество пассажиров с пропиской в КО.

Пользователь сверяет количество.
Если все цифры совпадают нажимает "Enter"

???
Постарается для каждого несовершеннолетнего 
найти однофамильца среди совершеннолетних пассажиров 
??? 

Также использует 2 образца (доковские файлы в формате *.docx)
в которые вносит данные по совершеннолетним и
несовершеннолетним (при наличии) пассажирам.

Полученный результат ДЛЯ КАЖДОГО ПАССАЖИРА записывает в отдельный файл
(для совершеннолетных и несовершеннолетних пассажиров - 
помечается в названии),
которые сохраняет в рабочей директории (папка расположения программы).

Дополнительно сообщает в интерактивном режиме об
успешности выполнения всех этапов алгоритма,
а также об успешном завершении в целом.

Название бланка для Взрослых - "Бланк_Взросл.docx"

"""

import pandas as pd
import os
from datetime import datetime as dt
import docx


#########################################
# ФУНКЦИИ!!!


#####

def my_check(my_file_name: str, excel_format: str):
    check_file = False
    raboch_papka = os.getcwd()
    i = 0
    while (check_file == False) and (i < 5):
        i += 1
        try:
            my_way = os.path.join(raboch_papka, my_file_name + excel_format)
            check_file = os.path.exists(my_way)
            print(check_file)
            print(my_way)
            if check_file:
                break
        except BaseException:
            print('Убедитесь, что в рабочей директории, расположенной по адресу: ' + raboch_papka + 
            '.\n Размещен xlsx-файл со следующим названием: ' + my_file_name)
            input()

#####
            
def dprt_point_is(sheetnames: list):
    
    for i in sheetnames:
        i = i.lower()
        if "казан" in i:
            depart_point = "г. Казань"
            break
        elif "новосиб" in i:
            depart_point = "г. Новосибирск"
            break
        else:
            depart_point = "г. Москва"
    
    print(f"Все пассажиры прибыли из {depart_point}?")
    print("Введите Y / N")
    change_depart_point = input()
    if change_depart_point.lower() == "n":
        depart_point = input("Изменить место отправления на (например, г. Казань): ")
    return(depart_point)

#####

def check_date_arrival():
    """
    Помогает проверить сегодняшнюю дату и
    принимает на вход дату прибытия, проверяет ее.
    В случае успешной проверки последней - возвращает ее.
    """
    weekday_list = [
    "Понедельник",
    "Вторник",
    "Среда",
    "Четверг",
    "Пятница",
    "Суббота",
    "Воскресенье"
    ]
    
    wrong_date = True
    while wrong_date:
        print("Сегодняшняя дата: ")

        today_date = dt.today().strftime('%d.%m.%Y')
        print(today_date)

        print("Сегодня")
        print(weekday_list[dt.weekday(dt.today())])
        print("Дата и день недели - правильные?")
        print("Введите Y / N")
        check_date = input()
        if check_date.lower() == "n":
            print("Проверьте дату и время на Вашем ПК!")
        else:
            print("OK!")
            wrong_date = False

    print()

    arrival_date = True
    while arrival_date:
        try:
            wrong_year_val = False
            wrong_month_val = False
            wrong_day_val = False
            print("Введите предполагаемую дату прибытия")        
            year = int(input('Введите год (ПОЛНОСТЬЮ!): '))
            if year < 0:
                wrong_year_val = True
            month = int(input('Введите месяц: '))
            if month < 1 or month > 12:
                wrong_month_val = True
            day = int(input('Введите день: '))
            if day < 1 or day > 31:
                wrong_day_val = True
            date_arrival = dt(year, month, day,0,0,0)
        except ValueError:
            print("Недопустимое значение даты. А именно")
            print()
            if wrong_year_val:
                print("год должен быть более 0!")
                del(wrong_year_val)
            print()
            if wrong_month_val:
                print("численное значение месяца должно быть в пределах от 1 до 12!")
                del(wrong_month_val)
            print()
            if wrong_day_val:
                print("число должно быть в пределах от 1 до 31!")
                del(wrong_day_val)   
            print()
            continue
        print()
        print("Вы ввели следующую дату прибытия: ", date_arrival.strftime('%d.%m.%Y'))
        cherez_skolk_dney_pribytie = date_arrival - dt.today()
        cherez_skolk_dney_pribytie = cherez_skolk_dney_pribytie.days
        if cherez_skolk_dney_pribytie < 0 or cherez_skolk_dney_pribytie > 5:
            print("Вы уверены? Дата прилета отличатеся от сегодняшней на:",
                 cherez_skolk_dney_pribytie, "дней.")
            print("Введите Y / N")
            my_input = input()
            if my_input.lower() == "n":
                print("Проверьте дату прибытия или время на Вашем ПК!")
            else:
                print("OK!")
                arrival_date = False
        else:
            arrival_date = False

    return(date_arrival)

#####

def check_FIO():
    print("Фамилия, Имя, Отчество находятся в одном столбце? ")
    print("Введите Y / N")
    FIO_vmeste = True
    my_input = input()
    if my_input.lower() == "n":
        print("Выбран ваиант ФИО - в разных столбцах")
        FIO_vmeste = False
    else:
        print("Выбран ваиант ФИО - в одном столбце")
    return(FIO_vmeste)

#####
  
def check_pass_amount(my_df, when_arrival):
    """
    Принимает на вход датафрэйм и дату прибытия
    подает на выход список из трех значений:
    количество детей,
    количество взрослых,
    общее количество пассажиров
    Возраст считается на дату прибытия
    Возвращает буллевый список взрослых пассажиров в таблице
    """
    child_amount = 0
    adult_amount = 0
    total_amount = 0
    my_bool = []
    total1 = sum(my_df.iloc[:,-1].fillna(0))-len(my_df.iloc[:,-1])
    total1 = int(total1)
    total2 = len(my_df.iloc[:,-1])
    if total1 == total2:
        total_amount = total1
    for i in my_df.iloc[:,-6]:
        if int((when_arrival - i).days/365) > 18:
            adult_amount += 1
            my_bool.append(True)
        else:
            child_amount += 1
            my_bool.append(False)
    print(f"Общее количество пассажиров: {total_amount}")
    print(f"Количество несовершеннолетних пассажиров: {child_amount}")
    print(f"Количество взрослых пассажиров: {adult_amount}")
    if child_amount + adult_amount != total_amount:
        print("Общее количество пассажиров не равно сумме взрослых и детей!")
    else:
        print("Сумма - ОК!")
    input()
    return(my_bool)
    
#####

def make_doc_magic_adult(data, FIO_tgthr, dep_re, arrvl_date):
    print("ВЗРОСЛЫЕ: ")
    print()
    if FIO_tgthr:
        FIO = list(data.iloc[:,1])
    else:
        FIO = []
        for i in range(len(data.iloc[:,1])):
            if ((len(str(data.iloc[i,3])) > 1) and 
            (str(data.iloc[i,3]).lower() != "nan")):
                FIO.append(data.iloc[i,1]+" "+
                data.iloc[i,2]+" "+
                str(data.iloc[i,3]))
            else:
                FIO.append(data.iloc[i,1]+" "+
                data.iloc[i,2])
    
    print(FIO)
    
    birth_date = []
    for x in data.iloc[:,-6]:
        birth_date.append(x.strftime('%d.%m.%Y'))
    print(birth_date)
    
    day_arrvl = arrvl_date.strftime('%d')
    
    month_arrvl = arrvl_date.strftime('%m')
    
    for j in range(len(data.iloc[:,1])):
        doc_number = data.iloc[j,-4]
        doc = docx.Document('Бланк_Взросл.docx')
        i = 0
        
        for paragraph in doc.paragraphs:
            i += 1
            if i == 14 or i == 30:
                paragraph.add_run("   "+day_arrvl+(" "*10)+month_arrvl)
            elif i == 18:
                paragraph.add_run("       "+FIO[j].upper()).bold = True
            elif i == 31 or i == 40:
                paragraph.add_run((" "*30)+FIO[j].upper()).bold = True
            elif i == 19:
                paragraph.add_run((" "*10)+birth_date[j])
            elif i == 22:
                paragraph.add_run((" "*30)+data.iloc[j,-3])
            elif i == 25 and str(doc_number).isdigit():
                if len(str(int(doc_number)).strip()) == 10:
                    doc_number = str(int(doc_number)).strip()[:4] + " " + str(int(doc_number))[4:]
                    paragraph.add_run("   "+doc_number+"   ПАСПОРТ").italic = True
                else:
                    paragraph.add_run("   "+str(int(doc_number))).italic = True
            elif i == 25:
                paragraph.add_run("   "+str(doc_number)).italic = True
            elif i == 26 and "nan" not in str(data.iloc[j,-5]).lower():
                paragraph.add_run("   "+str(data.iloc[j,-5]))
            elif i == 34:
                paragraph.add_run("   "+dep_re).bold = True


        doc.save("Взрослый_"+ FIO[j] +".docx")



#####



def make_doc_magic_child(data, FIO_tgthr, dep_re, arrvl_date):
    print("ДЕТИ:")
    print()
    if FIO_tgthr:
        FIO = list(data.iloc[:,1])
    else:
        FIO = []
        for i in range(len(data.iloc[:,1])):
            if ((len(str(data.iloc[i,3])) > 1) and 
            (str(data.iloc[i,3]).lower() != "nan")):
                FIO.append(data.iloc[i,1]+" "+
                data.iloc[i,2]+" "+
                str(data.iloc[i,3]))
            else:
                FIO.append(data.iloc[i,1]+" "+
                data.iloc[i,2])
    
    print(FIO)
    
    birth_date = []
    for x in data.iloc[:,-6]:
        birth_date.append(x.strftime('%d.%m.%Y'))
    print(birth_date)
    
    day_arrvl = arrvl_date.strftime('%d')
    
    month_arrvl = arrvl_date.strftime('%m')
    
    for j in range(len(data.iloc[:,1])):
        doc_number = data.iloc[j,-4]
        doc = docx.Document('Бланк_Дети.docx')
        i = 0
        
        for paragraph in doc.paragraphs:
            i += 1
            if i == 16 or i == 41:
                paragraph.add_run("   "+day_arrvl+(" "*10)+month_arrvl)
            elif i == 20:
                paragraph.add_run("       "+FIO[j].upper()).bold = True
            elif i == 42 or i == 54:
                paragraph.add_run((" "*30)+FIO[j].upper()).bold = True
            elif i == 21:
                paragraph.add_run((" "*10)+birth_date[j])
            elif i == 24:
                paragraph.add_run((" "*30)+data.iloc[j,-3])
            elif i == 27 and str(doc_number).isdigit():
                if len(str(int(doc_number)).strip()) == 10:
                    doc_number = str(int(doc_number)).strip()[:4] + " " + str(int(doc_number))[4:]
                    paragraph.add_run("   "+doc_number+"   ПАСПОРТ").italic = True
                else:
                    paragraph.add_run("   "+str(int(doc_number))).italic = True
            elif i == 27:
                paragraph.add_run("   "+str(doc_number)).italic = True
            elif i == 45:
                paragraph.add_run((" "*129)+dep_re).bold = True


        doc.save("Дети_"+ FIO[j] +".docx")
    print()
#####

#######################################################
         
try:
    file = input("Введите название файла (до точки): ")
    print("Для выбора формата .xls - введите 1")
    print("Для выбора формата .xlsx - введите 2")
    print("Для выбора формата .xlsb - введите 3")
    change_format = int(input("Введите 1 или 2 или 3: "))
    if change_format == 1:
        excel_format = '.xls'
    elif change_format == 2:
        excel_format = '.xlsx'
    elif change_format == 3:
        excel_format = '.xlsb'
    my_check(file, excel_format)
except BaseException:
    print("Ошибка ввода")
    input()
data = pd.ExcelFile(file+excel_format)
my_sheet_names = data.sheet_names
print("Названия листов в книге: ")
print(my_sheet_names)
nomer_lista = int(input("Введите номер листа: "))-1
df = data.parse(my_sheet_names[nomer_lista])
print()
#ФИО - в одном столбце?
FIO_vmeste = check_FIO()

df = df.iloc[1:,:]
if FIO_vmeste:
    df = df.iloc[:,:8]
else:
    df = df.iloc[:,:10]

#print(df)
departure_from = dprt_point_is(my_sheet_names)
print("Вы выбрали в качестве места отправления", departure_from)

print()

arrvl_date = check_date_arrival()
print()

# проверяем количество пассажиров в целом и по возрасту
l_bool = check_pass_amount(df, arrvl_date)

# ДФ по взрослым пассажирам
df_adult = df[l_bool]

# ДФ по несовершеннолетним пассажирам
if sum(l_bool) < len(df.iloc[:,0]):
    l_rev_bool = []
    for i in l_bool:
        if not(i):
            l_rev_bool.append(True)
        else:
            l_rev_bool.append(False)
    df_child = df[l_rev_bool]
    make_doc_magic_child(df_child, FIO_vmeste, departure_from, arrvl_date)

make_doc_magic_adult(df_adult, FIO_vmeste, departure_from, arrvl_date)
    
input("Нажмите клавишу 'Enter' для выхода")

